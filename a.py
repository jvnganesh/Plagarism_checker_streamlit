import streamlit as st
import nltk
import difflib
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

nltk.download('punkt')

def load_document(file):
    doc = Document(file)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]
    return paragraphs

def highlight_text(run, similarity):
    if similarity > 0.75:
        run.font.highlight_color = WD_COLOR_INDEX.RED
    elif similarity > 0.5:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def compare_documents(doc1, doc2):
    doc1_sentences = [sentence for para in doc1 for sentence in nltk.sent_tokenize(para)]
    doc2_sentences = [sentence for para in doc2 for sentence in nltk.sent_tokenize(para)]

    highlighted_paragraphs_doc1 = []
    highlighted_paragraphs_doc2 = []
    plagiarized_sentences_doc1 = 0
    plagiarized_sentences_doc2 = 0

    for sentence1 in doc1_sentences:
        similarity_scores = [difflib.SequenceMatcher(None, sentence1, sentence2).ratio() for sentence2 in doc2_sentences]
        if similarity_scores:
            max_similarity = max(similarity_scores)
            if max_similarity > 0.5:
                highlighted_paragraphs_doc1.append((sentence1, max_similarity))
                plagiarized_sentences_doc1 += 1

    for sentence2 in doc2_sentences:
        similarity_scores = [difflib.SequenceMatcher(None, sentence2, sentence1).ratio() for sentence1 in doc1_sentences]
        if similarity_scores:
            max_similarity = max(similarity_scores)
            if max_similarity > 0.5:
                highlighted_paragraphs_doc2.append((sentence2, max_similarity))
                plagiarized_sentences_doc2 += 1

    total_sentences_doc1 = len(doc1_sentences)
    total_sentences_doc2 = len(doc2_sentences)
    plagiarism_percentage_doc1 = (plagiarized_sentences_doc1 / total_sentences_doc1) * 100 if total_sentences_doc1 > 0 else 0
    plagiarism_percentage_doc2 = (plagiarized_sentences_doc2 / total_sentences_doc2) * 100 if total_sentences_doc2 > 0 else 0

    return highlighted_paragraphs_doc1, highlighted_paragraphs_doc2, plagiarism_percentage_doc1, plagiarism_percentage_doc2

def add_highlighted_content(doc, paragraphs, highlighted_paragraphs):
    for para in paragraphs:
        paragraph = doc.add_paragraph()
        sentences = nltk.sent_tokenize(para)
        for sentence in sentences:
            run = paragraph.add_run(sentence)
            for highlighted_sentence, similarity in highlighted_paragraphs:
                if sentence == highlighted_sentence:
                    highlight_text(run, similarity)

def display_pagewise_summary(paragraphs, highlighted_paragraphs):
    summary = {}
    page = 1
    char_count = 0
    page_char_limit = 2000  # This is an approximation, adjust as needed

    for para in highlighted_paragraphs:
        char_count += len(para[0])
        if char_count > page_char_limit:
            page += 1
            char_count = len(para[0])

        if page not in summary:
            summary[page] = []
        summary[page].append(para[0])

    return summary

st.title("Plagiarism Checker")

uploaded_file1 = st.file_uploader("Upload First Document", type=["docx"])
uploaded_file2 = st.file_uploader("Upload Second Document", type=["docx"])

if uploaded_file1 and uploaded_file2:
    doc1_paragraphs = load_document(uploaded_file1)
    doc2_paragraphs = load_document(uploaded_file2)

    highlighted_paragraphs_doc1, highlighted_paragraphs_doc2, plagiarism_percentage_doc1, plagiarism_percentage_doc2 = compare_documents(doc1_paragraphs, doc2_paragraphs)

    st.subheader("Highlighted Plagiarized Text")
    doc = Document()

    # Add plagiarism report for Document 1 on the first page
    report_paragraph = doc.add_paragraph()
    report_paragraph.add_run(f"Plagiarism Report for Document 1: {plagiarism_percentage_doc1:.2f}% of the content is plagiarized.\n\n").bold = True

    # Add Document 1 content with highlighted plagiarism
    add_highlighted_content(doc, doc1_paragraphs, highlighted_paragraphs_doc1)

    # Add a page break between the two documents
    doc.add_page_break()

    # Add plagiarism report for Document 2
    report_paragraph = doc.add_paragraph()
    report_paragraph.add_run(f"Plagiarism Report for Document 2: {plagiarism_percentage_doc2:.2f}% of the content is plagiarized.\n\n").bold = True

    # Add Document 2 content with highlighted plagiarism
    add_highlighted_content(doc, doc2_paragraphs, highlighted_paragraphs_doc2)

    doc.save("highlighted_plagiarism.docx")

    with open("highlighted_plagiarism.docx", "rb") as f:
        st.download_button("Download Highlighted Document", f, file_name="highlighted_plagiarism.docx")

    st.write(f"Plagiarism Report for Document 1: {plagiarism_percentage_doc1:.2f}% of the content is plagiarized.")
    st.write(f"Plagiarism Report for Document 2: {plagiarism_percentage_doc2:.2f}% of the content is plagiarized.")
    
    summary_doc1 = display_pagewise_summary(doc1_paragraphs, highlighted_paragraphs_doc1)
    summary_doc2 = display_pagewise_summary(doc2_paragraphs, highlighted_paragraphs_doc2)
    
    st.subheader("Page-wise Summary of Plagiarized Content for Document 1")
    for page, content in summary_doc1.items():
        st.write(f"Page {page}")
        for para in content:
            st.write(para)

    st.subheader("Page-wise Summary of Plagiarized Content for Document 2")
    for page, content in summary_doc2.items():
        st.write(f"Page {page}")
        for para in content:
            st.write(para)
